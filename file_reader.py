#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
统一文件读取工具 - file_reader.py
支持格式：PDF, DOCX, TXT, XLSX, CSV
自动编码检测，统一输出接口

使用方法：
    from file_reader import read_file, FileReader
    
    # 简单读取
    content = read_file("document.pdf")
    
    # 高级用法
    reader = FileReader()
    content = reader.read("document.docx")
"""

import os
import io
from pathlib import Path
from typing import Union, Optional, List, Dict, Any, Iterator
import logging

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class FileReaderError(Exception):
    """文件读取异常"""
    pass


class FileReader:
    """统一文件读取器"""
    
    # 支持的文件格式
    SUPPORTED_FORMATS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
        '.txt': 'text',
        '.csv': 'csv',
        '.xlsx': 'excel',
        '.xls': 'excel',
    }
    
    def __init__(self, encoding_fallbacks: List[str] = None):
        """
        初始化文件读取器
        
        Args:
            encoding_fallbacks: 编码检测失败时尝试的编码列表
        """
        self.encoding_fallbacks = encoding_fallbacks or [
            'utf-8', 'gbk', 'gb2312', 'gb18030', 'utf-16', 'utf-16-le', 'utf-16-be',
            'latin-1', 'cp1252', 'big5', 'shift_jis'
        ]
    
    def read(self, file_path: Union[str, Path], **options) -> Union[str, List[Dict], Any]:
        """
        读取文件内容
        
        Args:
            file_path: 文件路径
            **options: 额外选项
                - pdf_engine: PDF读取引擎 ('pdfplumber', 'pymupdf')
                - excel_sheet: Excel读取的sheet名称或索引
                - max_pages: PDF最大读取页数
                
        Returns:
            文件内容（根据格式返回不同类型）
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileReaderError(f"文件不存在: {file_path}")
        
        ext = file_path.suffix.lower()
        file_type = self.SUPPORTED_FORMATS.get(ext)
        
        if not file_type:
            # 尝试作为文本文件读取
            logger.warning(f"未知文件格式: {ext}，尝试作为文本文件读取")
            return self._read_text(file_path)
        
        try:
            if file_type == 'pdf':
                return self._read_pdf(file_path, **options)
            elif file_type == 'docx':
                return self._read_docx(file_path, **options)
            elif file_type == 'text':
                return self._read_text(file_path, **options)
            elif file_type == 'csv':
                return self._read_csv(file_path, **options)
            elif file_type == 'excel':
                return self._read_excel(file_path, **options)
            else:
                return self._read_text(file_path, **options)
        except Exception as e:
            logger.error(f"读取文件失败: {file_path}, 错误: {str(e)}")
            raise FileReaderError(f"读取文件失败: {str(e)}")
    
    def _detect_encoding(self, file_path: Path) -> str:
        """自动检测文件编码"""
        # 首先尝试使用 chardet
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(32768)  # 读取前32KB用于检测
                result = chardet.detect(raw_data)
                encoding = result.get('encoding', 'utf-8')
                confidence = result.get('confidence', 0)
                
                if encoding and confidence > 0.5:
                    logger.debug(f"检测到编码: {encoding} (置信度: {confidence:.2f})")
                    # 统一处理一些常见的编码别名
                    encoding = encoding.lower().replace('-', '_')
                    if encoding == 'gb2312':
                        encoding = 'gbk'  # GBK兼容GB2312
                    elif encoding == 'ascii':
                        encoding = 'utf-8'  # ASCII是UTF-8的子集
                    return encoding
        except ImportError:
            logger.debug("chardet未安装，使用备用检测方法")
        
        # 备用：尝试多种编码
        for encoding in self.encoding_fallbacks:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    f.read(1024)  # 尝试读取一部分
                    return encoding
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        # 默认使用UTF-8，并忽略错误
        logger.warning("无法确定编码，使用UTF-8（忽略错误）")
        return 'utf-8'
    
    def _read_text(self, file_path: Path, encoding: Optional[str] = None, 
                   errors: str = 'replace', **options) -> str:
        """读取文本文件"""
        if encoding is None:
            encoding = self._detect_encoding(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding, errors=errors) as f:
                content = f.read()
            logger.info(f"成功读取文本文件: {file_path}, 编码: {encoding}")
            return content
        except Exception as e:
            logger.error(f"读取文本文件失败: {e}")
            raise
    
    def _read_pdf(self, file_path: Path, pdf_engine: str = 'pdfplumber',
                  max_pages: Optional[int] = None, **options) -> str:
        """读取PDF文件"""
        if pdf_engine == 'pdfplumber':
            return self._read_pdf_pdfplumber(file_path, max_pages)
        else:
            return self._read_pdf_pymupdf(file_path, max_pages)
    
    def _read_pdf_pdfplumber(self, file_path: Path, max_pages: Optional[int] = None) -> str:
        """使用pdfplumber读取PDF"""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber未安装，尝试使用PyMuPDF")
            return self._read_pdf_pymupdf(file_path, max_pages)
        
        text_parts = []
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                pages_to_read = min(max_pages, total_pages) if max_pages else total_pages
                
                for i, page in enumerate(pdf.pages):
                    if i >= pages_to_read:
                        break
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            content = '\n\n'.join(text_parts)
            logger.info(f"成功读取PDF (pdfplumber): {file_path}, 共{pages_to_read}页")
            return content
        except Exception as e:
            logger.error(f"pdfplumber读取失败: {e}")
            raise
    
    def _read_pdf_pymupdf(self, file_path: Path, max_pages: Optional[int] = None) -> str:
        """使用PyMuPDF读取PDF"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise FileReaderError("PyMuPDF未安装，请运行: pip install PyMuPDF")
        
        text_parts = []
        try:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            pages_to_read = min(max_pages, total_pages) if max_pages else total_pages
            
            for i in range(pages_to_read):
                page = doc[i]
                text_parts.append(page.get_text())
            
            doc.close()
            content = '\n\n'.join(text_parts)
            logger.info(f"成功读取PDF (PyMuPDF): {file_path}, 共{pages_to_read}页")
            return content
        except Exception as e:
            logger.error(f"PyMuPDF读取失败: {e}")
            raise
    
    def _read_docx(self, file_path: Path, include_tables: bool = True, **options) -> str:
        """读取DOCX文件"""
        try:
            from docx import Document
        except ImportError:
            raise FileReaderError("python-docx未安装，请运行: pip install python-docx")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # 读取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 读取表格
            if include_tables:
                for table in doc.tables:
                    text_parts.append('\n[表格]')
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells]
                        text_parts.append(' | '.join(row_text))
                    text_parts.append('[表格结束]\n')
            
            content = '\n'.join(text_parts)
            logger.info(f"成功读取DOCX: {file_path}")
            return content
        except Exception as e:
            logger.error(f"读取DOCX失败: {e}")
            raise
    
    def _read_csv(self, file_path: Path, encoding: Optional[str] = None, 
                  **options) -> List[Dict]:
        """读取CSV文件"""
        try:
            import csv
        except ImportError:
            raise FileReaderError("csv模块不可用")
        
        if encoding is None:
            encoding = self._detect_encoding(file_path)
        
        try:
            with open(file_path, 'r', encoding=encoding, newline='') as f:
                # 尝试自动检测分隔符
                sample = f.read(8192)
                f.seek(0)
                
                sniffer = csv.Sniffer()
                try:
                    dialect = sniffer.sniff(sample)
                    delimiter = dialect.delimiter
                except:
                    delimiter = ','  # 默认逗号
                
                reader = csv.DictReader(f, delimiter=delimiter)
                rows = list(reader)
            
            logger.info(f"成功读取CSV: {file_path}, 共{len(rows)}行")
            return rows
        except Exception as e:
            logger.error(f"读取CSV失败: {e}")
            raise
    
    def _read_excel(self, file_path: Path, sheet_name=None, **options) -> Any:
        """读取Excel文件"""
        try:
            import pandas as pd
        except ImportError:
            raise FileReaderError("pandas未安装，请运行: pip install pandas openpyxl")
        
        try:
            # 使用pandas读取，自动处理多种格式
            df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
            
            # 返回DataFrame或字典列表
            if isinstance(df, dict):
                # 多个sheet的情况
                result = {name: sheet_df.to_dict('records') 
                         for name, sheet_df in df.items()}
                logger.info(f"成功读取Excel: {file_path}, 多个sheet")
                return result
            else:
                records = df.to_dict('records')
                logger.info(f"成功读取Excel: {file_path}, 共{len(records)}行")
                return records
        except Exception as e:
            logger.error(f"读取Excel失败: {e}")
            raise


# 便捷函数
def read_file(file_path: Union[str, Path], **options) -> Union[str, List[Dict], Any]:
    """
    便捷函数：读取文件内容
    
    Args:
        file_path: 文件路径
        **options: 读取选项
        
    Returns:
        文件内容
        
    Examples:
        >>> content = read_file("document.pdf")
        >>> content = read_file("document.docx")
        >>> content = read_file("data.xlsx", sheet_name="Sheet1")
    """
    reader = FileReader()
    return reader.read(file_path, **options)


def read_text(file_path: Union[str, Path], encoding: Optional[str] = None) -> str:
    """
    便捷函数：读取文本/PDF/DOCX内容为字符串
    
    Args:
        file_path: 文件路径
        encoding: 指定编码（仅对文本文件有效）
        
    Returns:
        文本内容
    """
    reader = FileReader()
    result = reader.read(file_path, encoding=encoding)
    
    # 如果返回的是列表（如Excel），转换为字符串
    if isinstance(result, list):
        return '\n'.join([str(row) for row in result])
    elif isinstance(result, dict):
        return str(result)
    return result


def batch_read(directory: Union[str, Path], pattern: str = "*", 
               recursive: bool = False, **options) -> Dict[str, Any]:
    """
    批量读取目录中的文件
    
    Args:
        directory: 目录路径
        pattern: 文件匹配模式，如 "*.pdf"
        recursive: 是否递归子目录
        **options: 读取选项
        
    Returns:
        文件名到内容的映射
        
    Examples:
        >>> results = batch_read("./docs", "*.pdf")
        >>> results = batch_read("./data", "*.txt", recursive=True)
    """
    directory = Path(directory)
    results = {}
    reader = FileReader()
    
    if recursive:
        files = list(directory.rglob(pattern))
    else:
        files = list(directory.glob(pattern))
    
    for file_path in files:
        if file_path.is_file():
            try:
                results[str(file_path)] = reader.read(file_path, **options)
            except Exception as e:
                logger.error(f"读取文件失败 {file_path}: {e}")
                results[str(file_path)] = f"[ERROR: {str(e)}]"
    
    logger.info(f"批量读取完成: {len(results)}/{len(files)} 个文件成功")
    return results


# 测试代码
if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        try:
            content = read_file(file_path)
            print(content[:2000] if len(str(content)) > 2000 else content)
            print(f"\n\n--- 读取完成 ---")
        except Exception as e:
            print(f"读取失败: {e}")
    else:
        print("""
统一文件读取工具

使用方法:
    python file_reader.py <文件路径>
    
或作为模块导入:
    from file_reader import read_file, FileReader
    content = read_file("document.pdf")
    
支持格式:
    - PDF (.pdf) - 自动使用pdfplumber或PyMuPDF
    - Word (.docx) - 使用python-docx
    - 文本 (.txt, .md, .json, etc.) - 自动编码检测
    - Excel (.xlsx, .xls) - 使用pandas+openpyxl
    - CSV (.csv) - 自动分隔符检测
        """)

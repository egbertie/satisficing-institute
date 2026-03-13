#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF和Word读取功能测试
"""

from pathlib import Path

print("=" * 60)
print("  PDF和Word读取功能测试")
print("=" * 60)
print()

# 检查依赖安装状态
print("[1] 检查依赖安装状态...")

deps = {
    'pdfplumber': 'PDF读取（推荐）',
    'fitz': 'PDF读取（高性能）',
    'docx': 'Word文档读取',
}

installed = []
missing = []

for module, desc in deps.items():
    try:
        __import__(module)
        print(f"  ✓ {module:15} - {desc}")
        installed.append(module)
    except ImportError:
        print(f"  ✗ {module:15} - {desc} (未安装)")
        missing.append(module)

if missing:
    print()
    print("安装缺失依赖:")
    print(f"  pip install pdfplumber PyMuPDF python-docx")
    print()

# 测试PDF读取
print()
print("[2] 测试PDF读取...")
try:
    import pdfplumber
    
    # 尝试读取系统中的PDF文件
    pdf_found = False
    for pdf_path in Path("/root").rglob("*.pdf"):
        if pdf_found:
            break
        pdf_found = True
        print(f"  找到PDF文件: {pdf_path}")
        
        with pdfplumber.open(pdf_path) as pdf:
            print(f"  页数: {len(pdf.pages)}")
            if len(pdf.pages) > 0:
                text = pdf.pages[0].extract_text()
                if text:
                    print(f"  首页内容预览: {text[:100]}...")
                else:
                    print("  首页无文本内容（可能是扫描件）")
        
    if not pdf_found:
        print("  未找到PDF测试文件，但pdfplumber已正确安装")
        
except Exception as e:
    print(f"  测试失败: {e}")

# 测试Word读取
print()
print("[3] 测试Word读取...")
try:
    from docx import Document
    print("  ✓ python-docx已正确安装")
    
    # 创建测试Word文档
    test_doc = Path("./test_files/test_document.docx")
    if not test_doc.exists():
        doc = Document()
        doc.add_heading('测试文档', 0)
        doc.add_paragraph('这是一个测试Word文档。')
        doc.add_paragraph('中文内容测试。')
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = '姓名'
        table.cell(0, 1).text = '年龄'
        table.cell(1, 0).text = '张三'
        table.cell(1, 1).text = '25'
        doc.save(test_doc)
        print(f"  ✓ 创建测试Word文档: {test_doc}")
    
    # 读取测试
    doc = Document(test_doc)
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    print(f"  ✓ 成功读取Word文档")
    print(f"    段落数: {len(doc.paragraphs)}")
    print(f"    表格数: {len(doc.tables)}")
    print(f"    内容预览: {' | '.join(text_parts[:3])}")
    
except Exception as e:
    print(f"  测试失败: {e}")

print()
print("=" * 60)
print("  测试完成")
print("=" * 60)
